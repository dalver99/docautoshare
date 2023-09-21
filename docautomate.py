from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load the document
doc = Document('inputdoc.docx')

# Search and edit the specific paragraph
for paragraph in doc.paragraphs:
    if paragraph.text == 'ADD RESULT:':
        paragraph.text = 'ADD RESULT: 8'

# Edit the first table, second row, second cell
table = doc.tables[0]
cell = table.cell(1, 1)
cell.text = 'A1'
paragraph = cell.paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
doc.save('output.docx')

#pip install python-docx
