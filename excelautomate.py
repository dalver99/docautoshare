from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlrd

# Load the workbook
wb = xlrd.open_workbook('input.xls')
sheet = wb.sheet_by_index(0)

# Load the document
doc = Document('inputdoc.docx')

# Fetch data from A2, A3 cell from excel
data_a2 = sheet.cell_value(1, 0)
data_a3 = sheet.cell_value(2, 0)

# Add this data
add_result = data_a2 + data_a3

# Search for the paragraph
for para in doc.paragraphs:
    if 'ADD RESULT:' in para.text:
        # Edit the paragraph
        para.text = f'ADD RESULT: {add_result}'

# Fetch data from A1 cell from excel
data_a1 = sheet.cell_value(0, 0)

# Write data to the first table in the document
table = doc.tables[0]
table.cell(0, 0).text = str(data_a1)

# Add a row at the bottom
row = table.add_row()
first_cell = row.cells[0]
first_cell.text = 'NEWROW'

# Save the document
doc.save('output2.docx')